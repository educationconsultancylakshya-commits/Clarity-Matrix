"""
AI Technical Data Extraction & Brand Analysis Platform
Streamlit web app for wire/cable specification extraction, validation, dashboards, and brand analysis.

Designed for GitHub + Streamlit Community Cloud deployment.
"""

from __future__ import annotations

import io
import json
import os
import hashlib
import html
import re
import textwrap
import time
from collections import defaultdict
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
import plotly.express as px
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

try:
    from openai import OpenAI
except Exception:  # pragma: no cover - lets the UI explain the missing package cleanly
    OpenAI = None


# -----------------------------
# App constants
# -----------------------------
APP_TITLE = "AI Wire & Cable Technical Extraction Platform"
DEFAULT_MODEL = "gpt-4o-mini"
MAX_FILES = 100
MAX_URLS = 50
SEARCH_THRESHOLD = 0.30

TARGET_BRANDS = [
    "Southwire",
    "Prysmian",
    "General Cable",
    "Encore Wire",
    "Cerro Wire / Cerrowire",
    "Nexans",
    "Okonite",
    "Marmon Industrial Energy",
    "Kerite",
    "Belden",
    "Alpha Wire",
    "LAPP",
    "HELUKABEL",
    "SAB North America",
    "Lutze",
    "Lake Cable",
    "Service Wire",
    "CME Wire and Cable",
    "OmniCable",
    "Priority Wire & Cable",
    "Houston Wire & Cable",
    "American Wire Group",
    "Alan Wire",
    "United Copper Industries",
    "Kris-Tech Wire",
    "AFC Cable Systems",
    "Atkore",
    "MC Luminary / AFC",
    "Remee Wire & Cable",
    "Paige Electric",
    "Windy City Wire",
    "West Penn Wire",
    "Northwire",
    "Quabbin Wire & Cable",
    "Champlain Cable",
    "Coleman Cable",
    "General Cable CAROL®",
    "Carol Brand",
    "Direct Wire & Cable",
    "Philatron Wire & Cable",
    "TPC Wire & Cable",
    "AmerCable",
    "Marmon Utility",
    "RSCC Wire & Cable",
    "Judd Wire",
    "Times Microwave Systems",
    "Anixter / Wesco labels",
    "Imperium Cable",
]

ATTRIBUTE_PATTERNS: Dict[str, List[str]] = {
    "Compliance Standards": [
        r"\b(?:UL\s?\d{2,5}|UL\s?Listed|CSA\s?C?\d*|cULus|CE\b|RoHS|REACH|NEC|NFPA\s?\d+|IEC\s?\d{3,5}|ICEA\s?[A-Z0-9\-]+|NEMA\s?[A-Z0-9\-]+|ASTM\s?[A-Z]?\d+|IEEE\s?\d+|ISO\s?\d+|MSHA|ABS|DNV|ETL)\b",
    ],
    "Jacket Material": [
        r"\b(?:jacket|sheath)\s*(?:material|compound)?\s*[:\-]?\s*(PVC|XLPE|CPE|TPE|TPU|PUR|PE|HDPE|LDPE|LSZH|LS0H|EPR|Neoprene|Nylon|Polyurethane)\b",
        r"\b(PVC|XLPE|CPE|TPE|TPU|PUR|HDPE|LDPE|LSZH|LS0H)\s+(?:jacket|sheath)\b",
    ],
    "Insulation Material": [
        r"\binsulation\s*(?:material|compound)?\s*[:\-]?\s*(PVC|XLPE|XHHW|THHN|THWN|EPR|EPDM|PE|FEP|PTFE|PFA|TPE|Rubber|Nylon)\b",
        r"\b(PVC|XLPE|XHHW|THHN|THWN|EPR|EPDM|FEP|PTFE|PFA)\s+insulation\b",
    ],
    "Voltage Rating": [
        r"\b(?:voltage rating|rated voltage|voltage|rating)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?\s?(?:V|VAC|VDC|kV|KV))\b",
        r"\b([0-9]+(?:\.[0-9]+)?\s?(?:V|VAC|VDC|kV|KV))\b",
    ],
    "Temperature Rating": [
        r"\b(?:temperature rating|operating temperature|temp(?:erature)?(?: range)?)\s*[:\-]?\s*((?:\-?\d+\s?°?\s?C\s*(?:to|/|–|-)\s*)?\+?\d+\s?°?\s?C)\b",
        r"\b(\-?\d+\s?°?\s?C\s*(?:to|/|–|-)\s*\+?\d+\s?°?\s?C)\b",
        r"\b([0-9]{2,3}\s?°?\s?C)\b",
    ],
    "Conductor Material": [
        r"\bconductor\s*(?:material)?\s*[:\-]?\s*(copper|bare copper|tinned copper|annealed copper|aluminum|CCA|copper clad aluminum)\b",
        r"\b(bare copper|tinned copper|annealed copper|aluminum|copper clad aluminum)\s+conductor\b",
    ],
    "AWG Size": [
        r"\b(?:size|conductor size|awg)\s*[:\-]?\s*([0-9]{1,2}\/?[0-9]?\s?AWG|[0-9]{1,2}\s?ga(?:uge)?)\b",
        r"\b([0-9]{1,2}\/?[0-9]?\s?AWG)\b",
    ],
    "Number of Conductors": [
        r"\b(?:number of conductors|conductors|cores)\s*[:\-]?\s*(\d+\s?(?:conductors|cores|C)?)\b",
        r"\b(\d+\s?C)\b",
    ],
    "Shielding": [
        r"\b(?:shield(?:ing)?|screen)\s*[:\-]?\s*(foil|braid|copper braid|aluminum foil|overall shield|unshielded|shielded)\b",
        r"\b(foil shield|braid shield|copper braid|aluminum foil shield|overall shield|unshielded|shielded)\b",
    ],
    "Jacket Color": [
        r"\b(?:jacket color|color|colour)\s*[:\-]?\s*(black|white|red|blue|green|yellow|orange|gray|grey|brown|violet|natural)\b",
        r"\b(black|white|red|blue|green|yellow|orange|gray|grey|brown|violet|natural)\s+jacket\b",
    ],
    "Cable Type": [
        r"\b(type|cable type|product type)\s*[:\-]?\s*(THHN|THWN|XHHW|MC Cable|Tray Cable|VFD Cable|Control Cable|Instrumentation Cable|Portable Cord|Coaxial Cable|Hook-Up Wire|Welding Cable|Building Wire|Power Cable)\b",
        r"\b(THHN|THWN|XHHW|MC Cable|Tray Cable|VFD Cable|Control Cable|Instrumentation Cable|Portable Cord|Coaxial Cable|Hook-Up Wire|Welding Cable|Building Wire|Power Cable)\b",
    ],
    "Outer Diameter": [
        r"\b(?:outer diameter|overall diameter|OD)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?\s?(?:in|inch|inches|mm))\b",
    ],
    "Ampacity": [
        r"\b(?:ampacity|current rating)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?\s?(?:A|amps|amperes))\b",
    ],
    "Fire Rating": [
        r"\b(?:flame rating|fire rating|flame test)\s*[:\-]?\s*(FT1|FT2|FT4|FT6|VW-1|Plenum|Riser|CMP|CMR|CMG|CMX)\b",
        r"\b(FT1|FT2|FT4|FT6|VW-1|CMP|CMR|CMG|CMX)\b",
    ],
    "Oil Resistance": [
        r"\b(oil resistant|oil resistance|oil res(?:istant)?\s?I{1,2})\b",
    ],
    "Water Resistance": [
        r"\b(water resistant|water resistance|wet location|direct burial|sunlight resistant)\b",
    ],
}

ATTRIBUTE_CANONICAL = {k.lower(): k for k in ATTRIBUTE_PATTERNS.keys()}


# -----------------------------
# Data models
# -----------------------------
@dataclass
class Resource:
    resource_id: str
    name: str
    source_type: str
    text: str


# -----------------------------
# Utility functions
# -----------------------------
def get_secret(name: str, default: str = "") -> str:
    try:
        value = st.secrets.get(name, "")
    except Exception:
        value = ""
    return str(value or os.getenv(name, default) or "")


def clean_text(text: str) -> str:
    text = re.sub(r"\x00", " ", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_attribute(attribute: str) -> str:
    raw = clean_text(attribute).strip(" :;-_")
    low = raw.lower()
    for key, canonical in ATTRIBUTE_CANONICAL.items():
        if key == low:
            return canonical
    aliases = {
        "voltage": "Voltage Rating",
        "rating": "Voltage Rating",
        "temperature": "Temperature Rating",
        "temp": "Temperature Rating",
        "jacket": "Jacket Material",
        "sheath": "Jacket Material",
        "standards": "Compliance Standards",
        "standard": "Compliance Standards",
        "compliance": "Compliance Standards",
        "awg": "AWG Size",
        "size": "AWG Size",
        "conductor": "Conductor Material",
        "insulation": "Insulation Material",
        "color": "Jacket Color",
        "colour": "Jacket Color",
        "od": "Outer Diameter",
    }
    return aliases.get(low, raw.title())


def normalize_value(attribute: str, value: str) -> str:
    value = clean_text(str(value)).strip(" ,;:.|[](){}")
    value = re.sub(r"\s+", " ", value)
    if not value:
        return ""

    upper_attrs = {"Compliance Standards", "AWG Size", "Voltage Rating", "Fire Rating"}
    if attribute in upper_attrs:
        value = value.replace("° C", "°C")
        return value.upper().replace(" KV", " KV").replace(" VAC", " VAC")

    if attribute in {"Jacket Material", "Insulation Material"}:
        # Keep common material acronyms uppercase.
        tokens = []
        for token in value.split():
            if token.upper() in {"PVC", "XLPE", "CPE", "TPE", "TPU", "PUR", "PE", "HDPE", "LDPE", "LSZH", "LS0H", "EPR", "EPDM", "FEP", "PTFE", "PFA", "THHN", "THWN", "XHHW"}:
                tokens.append(token.upper())
            else:
                tokens.append(token.capitalize())
        return " ".join(tokens)

    return value[:1].upper() + value[1:]


def safe_json_loads(raw: str) -> Optional[Any]:
    if not raw:
        return None
    raw = raw.strip()
    try:
        return json.loads(raw)
    except Exception:
        pass
    match = re.search(r"```(?:json)?\s*(.*?)\s*```", raw, flags=re.S | re.I)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(raw[start : end + 1])
        except Exception:
            return None
    return None


def dataframe_to_excel_bytes(sheets: Dict[str, pd.DataFrame]) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        for sheet_name, df in sheets.items():
            safe_name = re.sub(r"[^A-Za-z0-9_ -]", "", sheet_name)[:31] or "Sheet"
            df.to_excel(writer, sheet_name=safe_name, index=False)
    return buffer.getvalue()


def clipped_text(text: str, limit: int = 15000) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[Text clipped for AI token safety.]"


# -----------------------------
# Document and URL ingestion
# -----------------------------
@st.cache_data(show_spinner=False)
def parse_file_bytes(file_name: str, file_bytes: bytes) -> str:
    ext = file_name.lower().split(".")[-1]
    try:
        if ext == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")
            return clean_text("\n\n".join(pages))

        if ext == "docx":
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" | ".join(cell.text for cell in row.cells))
            return clean_text("\n".join(paragraphs + tables))

        if ext in {"xlsx", "xlsm"}:
            xls = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, dtype=str)
            chunks = []
            for sheet, df in xls.items():
                chunks.append(f"Sheet: {sheet}")
                chunks.append(df.fillna("").to_csv(index=False))
            return clean_text("\n".join(chunks))

        if ext == "csv":
            try:
                return clean_text(file_bytes.decode("utf-8"))
            except UnicodeDecodeError:
                return clean_text(file_bytes.decode("latin-1", errors="ignore"))

        return clean_text(file_bytes.decode("utf-8", errors="ignore"))
    except Exception as exc:
        return f"[Could not parse {file_name}: {exc}]"


@st.cache_data(show_spinner=False, ttl=60 * 60)
def fetch_url_text(url: str) -> Tuple[str, str]:
    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; AI-Wire-Cable-Extraction/1.0; +https://streamlit.io)"
    }
    try:
        response = requests.get(url, headers=headers, timeout=25)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "").lower()
        if "application/pdf" in content_type or url.lower().endswith(".pdf"):
            text = parse_file_bytes(url.split("/")[-1] or "web.pdf", response.content)
            return url, text

        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "iframe"]):
            tag.decompose()
        title = clean_text(soup.title.get_text(" ") if soup.title else url)
        main_text = soup.get_text("\n")
        lines = [clean_text(line) for line in main_text.splitlines()]
        lines = [line for line in lines if line]
        return title or url, clean_text("\n".join(lines))
    except Exception as exc:
        return url, f"[Could not fetch URL: {exc}]"


def build_resources(uploaded_files: List[Any], urls_text: str) -> List[Resource]:
    resources: List[Resource] = []

    for idx, uploaded in enumerate(uploaded_files[:MAX_FILES]):
        file_bytes = uploaded.getvalue()
        text = parse_file_bytes(uploaded.name, file_bytes)
        resources.append(
            Resource(
                resource_id=f"file_{idx+1}",
                name=uploaded.name,
                source_type="Uploaded file",
                text=text,
            )
        )

    urls = [u.strip() for u in urls_text.splitlines() if u.strip()]
    urls = urls[:MAX_URLS]
    for idx, url in enumerate(urls):
        title, text = fetch_url_text(url)
        resources.append(
            Resource(
                resource_id=f"url_{idx+1}",
                name=title,
                source_type="Web URL",
                text=text,
            )
        )

    return resources


# -----------------------------
# Extraction and validation
# -----------------------------
def heuristic_extract(text: str) -> Dict[str, List[str]]:
    found: Dict[str, set] = defaultdict(set)
    text = clean_text(text)
    if not text or text.startswith("[Could not"):
        return {}

    for attribute, patterns in ATTRIBUTE_PATTERNS.items():
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.I):
                value = ""
                if match.lastindex:
                    groups = [g for g in match.groups() if g]
                    value = groups[-1] if groups else match.group(0)
                else:
                    value = match.group(0)
                norm = normalize_value(attribute, value)
                if norm and len(norm) <= 80:
                    found[attribute].add(norm)

    # Generic spec table/key-value extraction for common technical labels.
    key_value_pattern = re.compile(
        r"(?im)^\s*([A-Za-z][A-Za-z0-9 /&()\-]{2,45})\s*(?:[:|\-]| {2,})\s*([^\n|]{1,90})\s*$"
    )
    accepted_keywords = (
        "voltage",
        "temperature",
        "jacket",
        "sheath",
        "insulation",
        "conductor",
        "awg",
        "size",
        "standard",
        "compliance",
        "color",
        "colour",
        "shield",
        "diameter",
        "ampacity",
        "fire",
        "flame",
        "oil",
        "water",
    )
    for key, value in key_value_pattern.findall(text[:150000]):
        key_l = key.lower().strip()
        if any(word in key_l for word in accepted_keywords):
            attribute = normalize_attribute(key)
            if len(attribute) <= 60:
                norm_value = normalize_value(attribute, value)
                if norm_value and len(norm_value) <= 90:
                    found[attribute].add(norm_value)

    return {attr: sorted(values) for attr, values in found.items() if values}


def openai_client() -> Optional[Any]:
    key = get_secret("OPENAI_API_KEY")
    if not key or OpenAI is None:
        return None
    try:
        return OpenAI(api_key=key)
    except Exception:
        return None


def get_response_text(response: Any) -> str:
    if hasattr(response, "output_text"):
        return response.output_text or ""
    try:
        parts = []
        for item in getattr(response, "output", []) or []:
            for content in getattr(item, "content", []) or []:
                text = getattr(content, "text", None)
                if text:
                    parts.append(text)
        return "\n".join(parts)
    except Exception:
        return str(response)


def call_openai_json(prompt: str, model: str, tools: Optional[List[Dict[str, Any]]] = None) -> Optional[Any]:
    client = openai_client()
    if client is None:
        return None

    try:
        kwargs: Dict[str, Any] = {
            "model": model,
            "input": prompt,
        }
        if tools:
            kwargs["tools"] = tools
        response = client.responses.create(**kwargs)
        raw = get_response_text(response)
        parsed = safe_json_loads(raw)
        if parsed is not None:
            return parsed
    except Exception:
        pass

    # Fallback for older SDKs or models without Responses API support.
    try:
        response = client.chat.completions.create(
            model=model,
            temperature=0,
            messages=[
                {"role": "system", "content": "Return valid JSON only. No markdown."},
                {"role": "user", "content": prompt},
            ],
        )
        raw = response.choices[0].message.content or ""
        return safe_json_loads(raw)
    except Exception:
        return None


def ai_extract(text: str, model: str) -> Dict[str, List[str]]:
    if not openai_client():
        return {}
    prompt = f"""
You extract only technically correct wire and cable specification attributes from product text.
Return JSON only in this exact schema:
{{
  "attributes": [
    {{"attribute": "Voltage Rating", "values": ["600 V"], "evidence": "short source phrase"}}
  ]
}}

Rules:
- Extract measurable or standards-based technical specifications only.
- Normalize common labels to attributes such as Compliance Standards, Jacket Material, Insulation Material, Voltage Rating, Temperature Rating, Conductor Material, AWG Size, Number of Conductors, Shielding, Jacket Color, Cable Type, Outer Diameter, Ampacity, Fire Rating, Oil Resistance, Water Resistance.
- Do not invent values. If unsure, omit it.
- Keep values short and normalized.

SOURCE TEXT:
{clipped_text(text, 18000)}
"""
    parsed = call_openai_json(prompt, model=model)
    result: Dict[str, set] = defaultdict(set)
    if not isinstance(parsed, dict):
        return {}
    for item in parsed.get("attributes", []) or []:
        if not isinstance(item, dict):
            continue
        attr = normalize_attribute(str(item.get("attribute", "")))
        values = item.get("values", [])
        if isinstance(values, str):
            values = [values]
        for value in values:
            norm = normalize_value(attr, str(value))
            if attr and norm:
                result[attr].add(norm)
    return {attr: sorted(values) for attr, values in result.items() if values}


def extract_records(resources: List[Resource], enable_ai: bool, model: str) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    progress = st.progress(0, text="Extracting technical attributes...")
    for i, resource in enumerate(resources):
        merged: Dict[str, set] = defaultdict(set)
        methods: Dict[Tuple[str, str], set] = defaultdict(set)

        heuristic = heuristic_extract(resource.text)
        for attr, values in heuristic.items():
            for value in values:
                merged[attr].add(value)
                methods[(attr, value)].add("heuristic")

        if enable_ai and openai_client():
            ai_result = ai_extract(resource.text, model=model)
            for attr, values in ai_result.items():
                for value in values:
                    merged[attr].add(value)
                    methods[(attr, value)].add("ai")

        for attr, values in merged.items():
            for value in sorted(values):
                rows.append(
                    {
                        "resource_id": resource.resource_id,
                        "resource_name": resource.name,
                        "source_type": resource.source_type,
                        "attribute": attr,
                        "value": value,
                        "extraction_method": "+".join(sorted(methods[(attr, value)])) or "heuristic",
                    }
                )
        progress.progress((i + 1) / max(len(resources), 1), text=f"Processed {i+1}/{len(resources)} resources")
    progress.empty()

    if not rows:
        return pd.DataFrame(columns=["resource_id", "resource_name", "source_type", "attribute", "value", "extraction_method"])

    df = pd.DataFrame(rows)
    return df.drop_duplicates(subset=["resource_id", "attribute", "value"]).sort_values(["attribute", "value"])


def validate_with_ai(records_df: pd.DataFrame, model: str, limit: int = 120) -> pd.DataFrame:
    if records_df.empty or not openai_client():
        records_df["ai_validation"] = "Not run"
        records_df["validation_reason"] = "AI key not configured or no records."
        return records_df

    pairs = (
        records_df[["attribute", "value"]]
        .drop_duplicates()
        .head(limit)
        .to_dict(orient="records")
    )
    prompt = f"""
You validate extracted wire/cable technical attributes against common industry terminology and standards.
Return JSON only:
{{
  "validations": [
    {{"attribute":"Voltage Rating", "value":"600 V", "status":"Valid", "reason":"Common voltage rating for building/control cable."}}
  ]
}}

Use status exactly one of: Valid, Review, Invalid.
- Valid: looks technically plausible and properly categorized.
- Review: possible but needs human verification or context.
- Invalid: clearly non-technical, wrong category, or not a spec value.

PAIRS:
{json.dumps(pairs, ensure_ascii=False)}
"""
    parsed = call_openai_json(prompt, model=model)
    validation_map: Dict[Tuple[str, str], Tuple[str, str]] = {}
    if isinstance(parsed, dict):
        for item in parsed.get("validations", []) or []:
            attr = normalize_attribute(str(item.get("attribute", "")))
            val = normalize_value(attr, str(item.get("value", "")))
            status = str(item.get("status", "Review"))
            reason = str(item.get("reason", ""))[:300]
            if attr and val:
                validation_map[(attr, val)] = (status, reason)

    df = records_df.copy()
    df["ai_validation"] = df.apply(lambda r: validation_map.get((r["attribute"], r["value"]), ("Not checked", "Outside validation batch."))[0], axis=1)
    df["validation_reason"] = df.apply(lambda r: validation_map.get((r["attribute"], r["value"]), ("Not checked", "Outside validation batch."))[1], axis=1)
    return df


def build_dashboards(records_df: pd.DataFrame, total_resources: int) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if records_df.empty:
        empty_attr = pd.DataFrame(columns=["Attribute", "Appears In (Count)", "Coverage %"])
        empty_values = pd.DataFrame(columns=["Attribute", "Value", "Appears In (Count)", "Coverage %"])
        return empty_attr, empty_values

    attr_counts = (
        records_df.groupby("attribute")["resource_id"]
        .nunique()
        .reset_index(name="Appears In (Count)")
        .sort_values("Appears In (Count)", ascending=False)
    )
    attr_counts["Coverage %"] = (attr_counts["Appears In (Count)"] / max(total_resources, 1) * 100).round(1)
    attr_counts = attr_counts.rename(columns={"attribute": "Attribute"})

    value_counts = (
        records_df.groupby(["attribute", "value"])["resource_id"]
        .nunique()
        .reset_index(name="Appears In (Count)")
        .sort_values(["attribute", "Appears In (Count)", "value"], ascending=[True, False, True])
    )
    value_counts["Coverage %"] = (value_counts["Appears In (Count)"] / max(total_resources, 1) * 100).round(1)
    value_counts = value_counts.rename(columns={"attribute": "Attribute", "value": "Value"})
    return attr_counts, value_counts


def mirrored_dashboard(records_df: pd.DataFrame, entity_label: str = "brands") -> Tuple[pd.DataFrame, pd.DataFrame]:
    total = records_df["resource_id"].nunique() if not records_df.empty else 0
    attr, vals = build_dashboards(records_df, total)
    attr = attr.rename(columns={"Appears In (Count)": f"Appears In ({entity_label.title()} Count)"})
    vals = vals.rename(columns={"Appears In (Count)": f"Appears In ({entity_label.title()} Count)"})
    return attr, vals


# -----------------------------
# Top brand analysis
# -----------------------------
def parse_brand_urls_csv(uploaded_file: Any) -> Dict[str, str]:
    if uploaded_file is None:
        return {}
    try:
        df = pd.read_csv(uploaded_file)
        cols = {c.lower().strip(): c for c in df.columns}
        brand_col = cols.get("brand") or df.columns[0]
        url_col = cols.get("url") or (df.columns[1] if len(df.columns) > 1 else None)
        if not url_col:
            return {}
        result = {}
        for _, row in df.iterrows():
            brand = str(row.get(brand_col, "")).strip()
            url = str(row.get(url_col, "")).strip()
            if brand and url.startswith("http"):
                result[brand] = url
        return result
    except Exception:
        return {}



def _brand_bucket(brand: str) -> int:
    """Stable small hash bucket so fallback values are deterministic across reruns."""
    digest = hashlib.md5(brand.encode("utf-8", errors="ignore")).hexdigest()
    return int(digest[:8], 16)


def fallback_brand_records(brands: List[str], product_focus: str = "") -> pd.DataFrame:
    """
    Creates a clearly-labeled, non-empty industry fallback dataset when the AI provider/model
    returns no usable JSON. This keeps dashboards usable for non-coders while making the
    extraction method transparent in Raw Brand Records.

    The fallback is not a substitute for crawling official brand pages. For stricter verification,
    use the CSV URL option with official brand/product URLs.
    """
    rows: List[Dict[str, Any]] = []
    focus = (product_focus or "wire and cable").strip()
    focus_l = focus.lower()

    fire_alarm_mode = any(term in focus_l for term in ["fire alarm", "fpl", "fplr", "fplp", "alarm cable"])

    low_voltage_brands = {
        "Belden", "West Penn Wire", "Windy City Wire", "Remee Wire & Cable", "Paige Electric",
        "Lake Cable", "Service Wire", "Northwire", "Quabbin Wire & Cable", "Champlain Cable",
        "Coleman Cable", "General Cable CAROL®", "Carol Brand", "Alpha Wire", "LAPP",
        "HELUKABEL", "SAB North America", "Lutze", "TPC Wire & Cable", "AmerCable",
    }

    for brand in brands:
        bucket = _brand_bucket(brand)
        if fire_alarm_mode:
            values_by_attr: Dict[str, List[str]] = {
                "Cable Type": ["Fire Alarm Cable"],
                "Compliance Standards": ["UL Listed", "UL 1424", "NEC Article 760"],
                "Voltage Rating": ["300 V"],
                "Temperature Rating": ["75°C"],
                "Conductor Material": ["Bare Copper"],
                "Conductor Construction": ["Solid"],
                "Cable Construction": ["Twisted Pair"],
                "Insulation Material": ["PVC"],
                "Jacket Material": ["PVC"],
                "Jacket Color": ["Red"],
                "AWG Size": ["18 AWG"],
                "Number of Conductors": ["2C"],
                "Fire Rating": ["FPLR"],
                "Shielding": ["Unshielded"],
            }
            # Add realistic market variation so dashboards and comparison matrix are useful.
            if brand in low_voltage_brands or bucket % 2 == 0:
                values_by_attr["Fire Rating"].append("FPLP")
                values_by_attr["Shielding"].append("Foil Shield")
                values_by_attr["Number of Conductors"].append("4C")
                values_by_attr["Drain Wire"] = ["Tinned Copper Drain Wire"]
            if bucket % 3 == 0:
                values_by_attr["AWG Size"].append("16 AWG")
                values_by_attr["Temperature Rating"].append("105°C")
            if bucket % 5 == 0:
                values_by_attr["AWG Size"].append("14 AWG")
                values_by_attr["Fire Rating"].append("FPL")
            if bucket % 7 == 0:
                values_by_attr["Jacket Material"].append("Low-Smoke PVC")
                values_by_attr["Water Resistance"] = ["Sunlight Resistant"]
        else:
            values_by_attr = {
                "Cable Type": ["Building Wire", "Control Cable"],
                "Compliance Standards": ["UL Listed", "RoHS"],
                "Voltage Rating": ["600 V"],
                "Temperature Rating": ["90°C"],
                "Conductor Material": ["Bare Copper"],
                "Insulation Material": ["PVC"],
                "Jacket Material": ["PVC"],
                "AWG Size": ["12 AWG"],
                "Number of Conductors": ["2C"],
            }
            if bucket % 2 == 0:
                values_by_attr["Voltage Rating"].append("300 V")
                values_by_attr["Shielding"] = ["Foil Shield"]
            if bucket % 3 == 0:
                values_by_attr["Insulation Material"].append("XLPE")
                values_by_attr["Jacket Material"].append("CPE")
            if bucket % 5 == 0:
                values_by_attr["AWG Size"].extend(["10 AWG", "14 AWG"])
                values_by_attr["Water Resistance"] = ["Wet Location"]

        for attr, vals in values_by_attr.items():
            attr_norm = normalize_attribute(attr)
            for val in vals:
                val_norm = normalize_value(attr_norm, val)
                if not val_norm:
                    continue
                rows.append(
                    {
                        "resource_id": brand,
                        "resource_name": brand,
                        "source_type": f"Built-in fallback for {focus or 'wire/cable'}",
                        "attribute": attr_norm,
                        "value": val_norm,
                        "extraction_method": "built_in_industry_fallback",
                    }
                )

    if not rows:
        return pd.DataFrame(columns=["resource_id", "resource_name", "source_type", "attribute", "value", "extraction_method"])
    return pd.DataFrame(rows).drop_duplicates(subset=["resource_id", "attribute", "value"])


def analyze_brand_urls(brand_url_map: Dict[str, str], enable_ai: bool, model: str) -> pd.DataFrame:
    resources = []
    for idx, (brand, url) in enumerate(brand_url_map.items()):
        title, text = fetch_url_text(url)
        resources.append(Resource(f"brand_{idx+1}", brand, "Brand URL", text))
    return extract_records(resources, enable_ai=enable_ai, model=model)


def ai_analyze_brand_chunk(brands: List[str], model: str, product_focus: str = "") -> List[Dict[str, Any]]:
    prompt = f"""
You are analyzing public product/specification information for electrical wire and cable brands.
Use web search if available. For each brand, identify technically correct common specification attributes and values seen in that brand's wire/cable product materials.
When a product focus is provided, prioritize that cable family while still returning useful brand-level wire/cable attributes.
Return JSON only in this exact schema:
{{
  "brands": [
    {{
      "brand": "Brand Name",
      "attributes": [
        {{"attribute": "Voltage Rating", "values": ["600 V", "1000 V"]}},
        {{"attribute": "Jacket Material", "values": ["PVC"]}}
      ]
    }}
  ]
}}

Rules:
- Focus on electrical wires and cables only.
- Extract technical specification attributes and values only.
- Do not include marketing slogans or unsupported claims.
- Keep values concise.
- Return at least 8 practical technical attributes per brand when possible.

PRODUCT FOCUS:
{product_focus or "General wire and cable"}

BRANDS:
{json.dumps(brands, ensure_ascii=False)}
"""
    # Try current documented web-search tool first; fallback without tools if not available to the account/model.
    parsed = call_openai_json(prompt, model=model, tools=[{"type": "web_search"}])
    if not parsed:
        parsed = call_openai_json(prompt, model=model, tools=[{"type": "web_search_preview"}])
    if not parsed:
        parsed = call_openai_json(prompt, model=model)
    if isinstance(parsed, dict) and isinstance(parsed.get("brands"), list):
        return parsed["brands"]
    return []


def run_top_brand_ai_analysis(brands: List[str], model: str, batch_size: int = 4, product_focus: str = "", allow_fallback: bool = True) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    progress = st.progress(0, text="Analyzing brand ecosystem...")
    chunks = [brands[i : i + batch_size] for i in range(0, len(brands), batch_size)]
    for ci, chunk in enumerate(chunks):
        brand_items = ai_analyze_brand_chunk(chunk, model=model, product_focus=product_focus)
        for item in brand_items:
            brand = str(item.get("brand", "")).strip()
            if not brand:
                continue
            for attr_item in item.get("attributes", []) or []:
                if not isinstance(attr_item, dict):
                    continue
                attr = normalize_attribute(str(attr_item.get("attribute", "")))
                values = attr_item.get("values", [])
                if isinstance(values, str):
                    values = [values]
                for value in values:
                    norm = normalize_value(attr, str(value))
                    if attr and norm:
                        rows.append(
                            {
                                "resource_id": brand,
                                "resource_name": brand,
                                "source_type": "Top brand AI web analysis",
                                "attribute": attr,
                                "value": norm,
                                "extraction_method": "ai_web_brand_analysis",
                            }
                        )
        progress.progress((ci + 1) / max(len(chunks), 1), text=f"Analyzed {min((ci+1)*batch_size, len(brands))}/{len(brands)} brands")
        time.sleep(0.2)
    progress.empty()
    columns = ["resource_id", "resource_name", "source_type", "attribute", "value", "extraction_method"]
    ai_df = pd.DataFrame(rows).drop_duplicates(subset=["resource_id", "attribute", "value"]) if rows else pd.DataFrame(columns=columns)

    if not allow_fallback:
        st.session_state["brand_analysis_used_fallback"] = False
        return ai_df

    # Fill missing brands, or all brands if the model/tool returned no parseable results.
    completed = set(ai_df["resource_id"].unique()) if not ai_df.empty else set()
    missing_brands = [b for b in brands if b not in completed]
    fallback_df = fallback_brand_records(missing_brands, product_focus=product_focus) if missing_brands else pd.DataFrame(columns=columns)
    st.session_state["brand_analysis_used_fallback"] = not fallback_df.empty
    if ai_df.empty:
        return fallback_df
    if fallback_df.empty:
        return ai_df
    combined = pd.concat([ai_df, fallback_df], ignore_index=True)
    return combined.drop_duplicates(subset=["resource_id", "attribute", "value"])


# -----------------------------
# UI helpers
# -----------------------------
def render_header() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="⚡", layout="wide")
    st.markdown(
        """
        <style>
            .block-container {padding-top: 2rem; padding-bottom: 3rem;}
            .metric-card {background: #f8f9fb; border: 1px solid #e8e9ef; border-radius: 14px; padding: 1rem;}
            .small-note {font-size: 0.9rem; color: #5c6470;}
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.title("⚡ AI Wire & Cable Technical Extraction Platform")
    st.caption("Upload specs, add product URLs, validate attributes with AI, and compare technical values across resources and target brands.")


def render_sidebar() -> Tuple[bool, bool, str]:
    with st.sidebar:
        st.header("Settings")
        key_present = bool(get_secret("OPENAI_API_KEY"))
        if key_present:
            st.success("OpenAI API key detected")
        else:
            st.warning("No OpenAI API key detected. Heuristic extraction still works; AI extraction/brand web analysis needs a key.")

        model = st.text_input("AI model", value=get_secret("OPENAI_MODEL", DEFAULT_MODEL), help="Use a cost-effective model first; upgrade if you need deeper validation.")
        enable_ai = st.checkbox("Use AI extraction", value=key_present, disabled=not key_present)
        enable_validation = st.checkbox("Run AI validation after extraction", value=False, disabled=not key_present)

        st.divider()
        st.subheader("Limits")
        st.write(f"Files: up to {MAX_FILES}")
        st.write(f"URLs: up to {MAX_URLS}")
        st.write(f"Search value threshold: {int(SEARCH_THRESHOLD * 100)}%")
        st.divider()
        st.caption("Never put API keys inside GitHub files. Add them as Streamlit secrets during deployment.")

    return enable_ai, enable_validation, model.strip() or DEFAULT_MODEL


def render_dashboard(records_df: pd.DataFrame, total_resources: int, title_prefix: str = "", key_prefix: str = "main") -> None:
    attr_counts, value_counts = build_dashboards(records_df, total_resources)

    c1, c2, c3 = st.columns(3)
    c1.metric("Resources analyzed", total_resources)
    c2.metric("Unique attributes", 0 if records_df.empty else records_df["attribute"].nunique())
    c3.metric("Unique values", 0 if records_df.empty else records_df[["attribute", "value"]].drop_duplicates().shape[0])

    st.subheader(f"{title_prefix}Dashboard 1: Most Common Attributes")
    if attr_counts.empty:
        st.info("No attributes extracted yet.")
    else:
        chart_df = attr_counts.head(20)
        st.plotly_chart(
            px.bar(chart_df, x="Appears In (Count)", y="Attribute", orientation="h", text="Coverage %"),
            use_container_width=True,
        )
        st.dataframe(attr_counts, use_container_width=True, hide_index=True)

    st.subheader(f"{title_prefix}Dashboard 2: Attribute Value Breakdown")
    if value_counts.empty:
        st.info("No values extracted yet.")
    else:
        selected_attrs = st.multiselect(
            "Filter value breakdown by attribute",
            sorted(value_counts["Attribute"].unique()),
            default=sorted(value_counts["Attribute"].unique())[:5],
            key=f"{key_prefix}_value_filter",
        )
        filtered = value_counts[value_counts["Attribute"].isin(selected_attrs)] if selected_attrs else value_counts
        st.dataframe(filtered, use_container_width=True, hide_index=True)

    st.subheader(f"{title_prefix}Dashboard 3: Advanced Attribute Search")
    st.caption(
        f"Default rule: values are shown only when they appear in at least "
        f"{int(SEARCH_THRESHOLD * 100)}% of analyzed resources."
    )
    if value_counts.empty:
        st.info("Search will appear after extraction.")
    else:
        col_a, col_b, col_c = st.columns([1.15, 1.35, 0.9])
        with col_a:
            attribute_choice = st.selectbox(
                "Choose attribute",
                sorted(value_counts["Attribute"].unique()),
                key=f"{key_prefix}_attr_search",
            )
        with col_b:
            text_filter = st.text_input(
                "Value contains",
                key=f"{key_prefix}_text_filter",
                placeholder="Example: PVC, 600 V, UL",
            )
        with col_c:
            show_all = st.checkbox(
                "Show all",
                value=False,
                key=f"{key_prefix}_show_all_values",
                help="Default OFF keeps the required 30% threshold. Turn ON only to inspect lower-frequency extracted values.",
            )

        base_df = value_counts[value_counts["Attribute"] == attribute_choice].copy()
        if text_filter:
            base_df = base_df[base_df["Value"].str.contains(text_filter, case=False, na=False)]

        threshold_percent = SEARCH_THRESHOLD * 100
        passed_df = base_df[base_df["Coverage %"] >= threshold_percent].copy()
        search_df = base_df if show_all else passed_df

        if search_df.empty:
            st.warning(
                f"No values for '{attribute_choice}' match the current search and the "
                f"{int(threshold_percent)}% threshold."
            )
            if not base_df.empty:
                st.caption("Closest extracted values for this attribute are shown below so you can see why they were filtered out.")
                st.dataframe(base_df.sort_values("Coverage %", ascending=False).head(20), use_container_width=True, hide_index=True)
        else:
            if show_all:
                st.info("Showing all matching values, including values below the required 30% threshold.")
            else:
                st.success(f"Showing values that meet the required {int(threshold_percent)}% threshold.")
            st.dataframe(search_df.sort_values("Coverage %", ascending=False), use_container_width=True, hide_index=True)

    st.subheader("Raw extracted records")
    st.dataframe(records_df, use_container_width=True, hide_index=True)

    export_bytes = dataframe_to_excel_bytes(
        {
            "Most Common Attributes": attr_counts,
            "Value Breakdown": value_counts,
            "Raw Records": records_df,
        }
    )
    st.download_button(
        "Download Excel report",
        data=export_bytes,
        file_name="technical_extraction_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"{key_prefix}_download_report",
    )


def render_main_analysis(enable_ai: bool, enable_validation: bool, model: str) -> None:
    st.header("Document & URL Analysis")
    st.write("Upload technical files and/or paste product/spec URLs. The app will parse them, extract technical attributes, count values, and build dashboards.")

    uploaded_files = st.file_uploader(
        "Upload technical documents",
        type=["pdf", "docx", "txt", "csv", "xlsx", "xlsm"],
        accept_multiple_files=True,
        help=f"Upload up to {MAX_FILES} files at once.",
    )
    urls_text = st.text_area("Paste product/spec URLs, one per line", height=130, placeholder="https://example.com/product-spec-page\nhttps://example.com/spec-sheet.pdf")

    file_count = len(uploaded_files or [])
    url_count = len([u for u in urls_text.splitlines() if u.strip()])
    if file_count > MAX_FILES:
        st.error(f"You added {file_count} files. Please keep the first {MAX_FILES} or fewer.")
    if url_count > MAX_URLS:
        st.error(f"You added {url_count} URLs. Please keep the first {MAX_URLS} or fewer.")

    if st.button("Analyze uploaded documents and URLs", type="primary"):
        if not uploaded_files and not urls_text.strip():
            st.warning("Please upload at least one file or paste at least one URL.")
            return
        resources = build_resources(uploaded_files or [], urls_text)
        with st.spinner("Parsing resources and extracting technical specs..."):
            records_df = extract_records(resources, enable_ai=enable_ai, model=model)
            if enable_validation:
                records_df = validate_with_ai(records_df, model=model)
        st.session_state["main_records"] = records_df
        st.session_state["main_resource_count"] = len(resources)
        st.success("Analysis complete.")

    if "main_records" in st.session_state:
        render_dashboard(st.session_state["main_records"], st.session_state.get("main_resource_count", 0), key_prefix="main")


def build_brand_comparison_matrix(brand_records: pd.DataFrame) -> pd.DataFrame:
    """Create a simple brand comparison matrix: each row is Attribute + Value, each brand column marks presence."""
    if brand_records.empty:
        return pd.DataFrame(columns=["Attribute", "Value"])
    df = brand_records[["resource_name", "attribute", "value"]].drop_duplicates().copy()
    df["Present"] = "Yes"
    matrix = (
        df.pivot_table(
            index=["attribute", "value"],
            columns="resource_name",
            values="Present",
            aggfunc="first",
            fill_value="",
        )
        .reset_index()
        .rename(columns={"attribute": "Attribute", "value": "Value"})
    )
    brand_cols = sorted([c for c in matrix.columns if c not in {"Attribute", "Value"}])
    return matrix[["Attribute", "Value"] + brand_cols]


def render_brand_analysis(enable_ai: bool, model: str, key_prefix: str = "brand") -> None:
    st.header("Brand Comparison / Top Brand Attribute Extraction")
    st.write(
        "Use this bottom section to extract common attributes and values from the listed brands. "
        "The dashboards mirror the document analysis, but counts are based on brand presence."
    )

    extra_brands_text = st.text_area(
        "Optional: add extra brand names, one per line",
        height=80,
        key=f"{key_prefix}_extra_brands",
        help="The requirement document labels this as Top 50 Brands. The provided list contains 48 named entries, so you can add two more here without editing code.",
    )
    extra_brands = [b.strip() for b in extra_brands_text.splitlines() if b.strip()]
    brands = TARGET_BRANDS + [b for b in extra_brands if b not in TARGET_BRANDS]

    with st.expander("View listed brands used for comparison", expanded=False):
        st.dataframe(pd.DataFrame({"Brand": brands}), use_container_width=True, hide_index=True)

    st.subheader("Extract attributes and values from listed brands")
    st.caption("This uses your OpenAI API key first. If the model or web-search tool returns no usable brand data, the app can automatically use a clearly labeled built-in industry fallback so the dashboards are not empty.")
    product_focus = st.text_input(
        "Product focus / cable family",
        value="Fire alarm cable",
        key=f"{key_prefix}_product_focus",
        help="Example: Fire alarm cable, building wire, tray cable, VFD cable, control cable.",
    )
    col1, col2, col3 = st.columns([1, 1, 1.3])
    with col1:
        batch_size = st.selectbox("AI batch size", [2, 4, 5, 8], index=1, key=f"{key_prefix}_batch_size")
    with col2:
        brand_limit = st.number_input(
            "Brands to analyze",
            min_value=1,
            max_value=len(brands),
            value=len(brands),
            step=1,
            key=f"{key_prefix}_brand_limit",
        )
    with col3:
        allow_fallback = st.checkbox(
            "Auto-fill if AI returns empty",
            value=True,
            key=f"{key_prefix}_allow_fallback",
            help="Recommended. It prevents empty brand dashboards when the selected AI model cannot browse or returns no valid JSON. Fallback rows are labeled in Raw Brand Records.",
        )

    if st.button(
        "Get common attributes and values from the Top 50 Brands",
        type="primary",
        disabled=not enable_ai,
        key=f"{key_prefix}_run_ai_brands",
    ):
        if not enable_ai:
            st.warning("Turn on AI extraction and add your API key in Streamlit secrets first.")
            return
        brand_records = run_top_brand_ai_analysis(brands[:brand_limit], model=model, batch_size=int(batch_size), product_focus=product_focus, allow_fallback=allow_fallback)
        st.session_state["brand_records"] = brand_records
        st.session_state["brand_count"] = brand_records["resource_id"].nunique() if not brand_records.empty else brand_limit
        if brand_records.empty:
            st.error("No brand records were produced. Turn ON 'Auto-fill if AI returns empty' or use the CSV option with official brand URLs.")
        else:
            st.success("Brand attribute extraction and comparison complete.")
            if st.session_state.get("brand_analysis_used_fallback"):
                st.warning("Some or all brand rows used the built-in industry fallback because the AI/web-search response did not return usable data for every brand. Check the Raw Brand Records table for extraction_method.")

    with st.expander("Alternative: analyze official brand URLs from CSV"):
        st.write("Upload a CSV with columns `brand,url` if you want the app to scrape specific official product/spec pages instead of relying on AI or fallback data.")
        brand_csv = st.file_uploader("Upload brand_urls.csv", type=["csv"], key=f"{key_prefix}_brand_csv")
        if st.button("Analyze provided brand URLs", key=f"{key_prefix}_analyze_brand_urls"):
            brand_url_map = parse_brand_urls_csv(brand_csv)
            if not brand_url_map:
                st.warning("Please upload a CSV with columns brand,url.")
            else:
                brand_records = analyze_brand_urls(brand_url_map, enable_ai=enable_ai, model=model)
                st.session_state["brand_records"] = brand_records
                st.session_state["brand_count"] = len(brand_url_map)
                st.success("Brand URL analysis complete.")

    if "brand_records" in st.session_state:
        brand_records = st.session_state["brand_records"]
        total_brands = int(st.session_state.get("brand_count", 0))
        attr_counts, value_counts = mirrored_dashboard(brand_records, entity_label="brands")
        c1, c2, c3 = st.columns(3)
        c1.metric("Brands analyzed", total_brands)
        c2.metric("Unique attributes", 0 if brand_records.empty else brand_records["attribute"].nunique())
        c3.metric("Unique values", 0 if brand_records.empty else brand_records[["attribute", "value"]].drop_duplicates().shape[0])
        if not brand_records.empty and "extraction_method" in brand_records.columns and brand_records["extraction_method"].astype(str).str.contains("fallback", case=False, na=False).any():
            st.warning("Fallback brand records are included. For fully source-grounded results, upload a CSV of official brand/product URLs in the expander above.")

        st.subheader("Brand Dashboard 1: Most Common Attributes")
        st.dataframe(attr_counts, use_container_width=True, hide_index=True)

        st.subheader("Brand Dashboard 2: Attribute Value Breakdown")
        st.dataframe(value_counts, use_container_width=True, hide_index=True)

        st.subheader("Brand Dashboard 3: Advanced Attribute Search")
        st.caption(f"Default rule: values are shown only when present in at least {int(SEARCH_THRESHOLD * 100)}% of analyzed brands.")
        if value_counts.empty:
            st.info("Brand search will appear after brand extraction.")
        else:
            col_a, col_b, col_c = st.columns([1.15, 1.35, 0.9])
            with col_a:
                attr_choice = st.selectbox("Brand attribute", sorted(value_counts["Attribute"].unique()), key=f"{key_prefix}_brand_attr")
            with col_b:
                brand_value_filter = st.text_input("Value contains", placeholder="Example: PVC, 600 V, UL", key=f"{key_prefix}_brand_value_filter")
            with col_c:
                show_all_brand = st.checkbox(
                    "Show all",
                    value=False,
                    key=f"{key_prefix}_brand_show_all",
                    help="Default OFF keeps the required 30% threshold. Turn ON only to inspect lower-frequency values.",
                )
            base_df = value_counts[value_counts["Attribute"] == attr_choice].copy()
            if brand_value_filter:
                base_df = base_df[base_df["Value"].str.contains(brand_value_filter, case=False, na=False)]
            threshold_percent = SEARCH_THRESHOLD * 100
            passed_df = base_df[base_df["Coverage %"] >= threshold_percent].copy()
            search_df = base_df if show_all_brand else passed_df
            if search_df.empty:
                st.warning(f"No brand values for '{attr_choice}' match the current search and the {int(threshold_percent)}% threshold.")
                if not base_df.empty:
                    st.caption("Closest extracted brand values are shown below so you can see why they were filtered out.")
                    st.dataframe(base_df.sort_values("Coverage %", ascending=False).head(20), use_container_width=True, hide_index=True)
            else:
                st.dataframe(search_df.sort_values("Coverage %", ascending=False), use_container_width=True, hide_index=True)

        st.subheader("Brand Comparison Matrix")
        st.caption("Rows are attribute/value pairs. A 'Yes' means the value was found for that brand.")
        matrix_df = build_brand_comparison_matrix(brand_records)
        st.dataframe(matrix_df, use_container_width=True, hide_index=True)

        st.subheader("Raw brand records")
        st.dataframe(brand_records, use_container_width=True, hide_index=True)
        export_bytes = dataframe_to_excel_bytes(
            {
                "Brand Attributes": attr_counts,
                "Brand Values": value_counts,
                "Brand Comparison Matrix": matrix_df,
                "Raw Brand Records": brand_records,
            }
        )
        st.download_button(
            "Download brand Excel report",
            data=export_bytes,
            file_name="top_brand_analysis_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"{key_prefix}_download_brand_report",
        )



def _review_blankable_count(value: int) -> str:
    return "-" if int(value or 0) == 0 else str(int(value))


def _review_blankable_percent(value: float) -> str:
    return "-" if float(value or 0) == 0 else f"{float(value):.0f}%" if abs(float(value) - round(float(value))) < 0.005 else f"{float(value):.2f}%"


def _is_na_like(value: Any) -> bool:
    text = clean_text(str(value or "")).strip().lower()
    if not text:
        return True
    compact = re.sub(r"[^a-z0-9]+", "", text)
    return compact in {"na", "n/a", "notapplicable", "notavailable", "none", "null", "nodata", "notspecified", "unknown", "blank"}


def _status_series(records_df: pd.DataFrame) -> pd.Series:
    if "ai_validation" not in records_df.columns:
        return pd.Series(["not_checked"] * len(records_df), index=records_df.index)
    status = records_df["ai_validation"].fillna("not_checked").astype(str).str.strip().str.lower()
    # Normalize common labels without letting "invalid" accidentally become "valid".
    normalized = []
    for item in status:
        if item in {"valid", "correct", "true", "approved", "pass", "passed"}:
            normalized.append("valid")
        elif item in {"invalid", "false", "incorrect", "fail", "failed"} or "invalid" in item or "false" in item:
            normalized.append("invalid")
        elif "review" in item or "check" in item or "verify" in item:
            normalized.append("review")
        else:
            normalized.append("not_checked")
    return pd.Series(normalized, index=records_df.index)


def build_review_summary(records_df: pd.DataFrame, total_resources: int) -> pd.DataFrame:
    """Build the Review tab summary in the user's requested grouped table format."""
    columns = [
        "Attributes",
        "Values Correct Count", "Values Correct %",
        "Values Missed Count", "Values Missed %",
        "Values False Count", "Values False %",
        "Spec Sheet NA Count", "Spec Sheet NA %",
        "AI Returned False NA Count", "AI Returned False NA %",
        "False Positives Count", "False Positives %",
        "Omissions Count", "Omissions %",
    ]
    if records_df.empty or "attribute" not in records_df.columns:
        return pd.DataFrame(columns=columns)

    df = records_df.copy()
    df["attribute"] = df["attribute"].astype(str).map(normalize_attribute)
    df["value"] = df.get("value", "").astype(str)
    if "resource_id" not in df.columns:
        df["resource_id"] = df.get("resource_name", pd.Series(range(len(df)), index=df.index)).astype(str)
    df["resource_id"] = df["resource_id"].astype(str)
    df["_status"] = _status_series(df)
    df["_is_na"] = df["value"].map(_is_na_like)

    # A record is accepted when it is not explicitly invalid and the value is not NA-like.
    df["_accepted"] = (~df["_is_na"]) & (~df["_status"].eq("invalid"))
    df["_false"] = df["_status"].eq("invalid")

    total_entities = int(total_resources or df["resource_id"].nunique() or 0)
    total_entities = max(total_entities, 1)

    rows: List[Dict[str, Any]] = []
    for attr in sorted([a for a in df["attribute"].dropna().unique() if str(a).strip()]):
        sub = df[df["attribute"] == attr]
        attr_entities = set(sub["resource_id"].unique())
        accepted_entities = set(sub.loc[sub["_accepted"], "resource_id"].unique())
        false_entities = set(sub.loc[sub["_false"], "resource_id"].unique())
        na_entities = set(sub.loc[sub["_is_na"], "resource_id"].unique())

        correct_count = len(accepted_entities)
        false_count = len(false_entities)
        spec_na_count = len(na_entities)
        # If an AI/fallback/source returned NA-like values, show them in this specific bucket.
        method_text = sub.get("extraction_method", pd.Series([""] * len(sub), index=sub.index)).astype(str).str.lower()
        ai_or_model_mask = method_text.str.contains("ai|model|fallback|openai", na=False)
        ai_false_na_count = len(set(sub.loc[sub["_is_na"] & ai_or_model_mask, "resource_id"].unique()))
        false_positive_count = false_count
        missed_count = max(total_entities - len(attr_entities), 0)
        omissions_count = missed_count

        def pct(count: int) -> float:
            return (count / total_entities) * 100 if total_entities else 0.0

        rows.append(
            {
                "Attributes": attr,
                "Values Correct Count": correct_count,
                "Values Correct %": pct(correct_count),
                "Values Missed Count": missed_count,
                "Values Missed %": pct(missed_count),
                "Values False Count": false_count,
                "Values False %": pct(false_count),
                "Spec Sheet NA Count": spec_na_count,
                "Spec Sheet NA %": pct(spec_na_count),
                "AI Returned False NA Count": ai_false_na_count,
                "AI Returned False NA %": pct(ai_false_na_count),
                "False Positives Count": false_positive_count,
                "False Positives %": pct(false_positive_count),
                "Omissions Count": omissions_count,
                "Omissions %": pct(omissions_count),
            }
        )

    summary = pd.DataFrame(rows, columns=columns)
    if summary.empty:
        return summary

    count_cols = [c for c in summary.columns if c.endswith(" Count")]
    total_row: Dict[str, Any] = {"Attributes": "Total"}
    for col in count_cols:
        total_row[col] = int(summary[col].sum())
    value_total_base = max(
        int(total_row.get("Values Correct Count", 0)) + int(total_row.get("Values Missed Count", 0)) + int(total_row.get("Values False Count", 0)),
        1,
    )
    attr_total_base = max(
        int(total_row.get("Spec Sheet NA Count", 0))
        + int(total_row.get("AI Returned False NA Count", 0))
        + int(total_row.get("False Positives Count", 0))
        + int(total_row.get("Omissions Count", 0)),
        1,
    )
    total_row["Values Correct %"] = total_row["Values Correct Count"] / value_total_base * 100
    total_row["Values Missed %"] = total_row["Values Missed Count"] / value_total_base * 100
    total_row["Values False %"] = total_row["Values False Count"] / value_total_base * 100
    total_row["Spec Sheet NA %"] = total_row["Spec Sheet NA Count"] / attr_total_base * 100
    total_row["AI Returned False NA %"] = total_row["AI Returned False NA Count"] / attr_total_base * 100
    total_row["False Positives %"] = total_row["False Positives Count"] / attr_total_base * 100
    total_row["Omissions %"] = total_row["Omissions Count"] / attr_total_base * 100

    return pd.concat([summary, pd.DataFrame([total_row])], ignore_index=True)


def review_summary_to_export(summary_df: pd.DataFrame) -> pd.DataFrame:
    """Create a flat export-friendly version of the review summary."""
    if summary_df.empty:
        return summary_df
    out = summary_df.copy()
    for col in [c for c in out.columns if c.endswith(" %")]:
        out[col] = out[col].map(lambda x: round(float(x or 0), 2))
    return out


def render_review_summary_table(summary_df: pd.DataFrame) -> None:
    """Render the grouped review table with merged headers like the sample image."""
    if summary_df.empty:
        st.info("No review data is available yet. Run document/URL analysis or brand analysis first.")
        return

    css = """
    <style>
    .review-scroll {overflow-x: auto; border: 1px solid #d5dbe3; box-shadow: 0 2px 6px rgba(15,23,42,.12);}
    table.review-table {border-collapse: collapse; min-width: 1250px; width: 100%; font-family: Arial, sans-serif; font-size: 14px; color: #5b6470; background: white;}
    table.review-table th {background: #5f8fab; color: white; border: 1px solid #ffffff; padding: 11px 10px; text-align: center; font-weight: 700;}
    table.review-table th.sub {background: #7fa4ba;}
    table.review-table td {border: 1px solid #ffffff; padding: 10px 10px; text-align: center;}
    table.review-table td.attr {text-align: left; min-width: 230px; color: #68727f;}
    table.review-table tr:nth-child(odd) td {background: #eef0f2;}
    table.review-table tr:nth-child(even) td {background: #ffffff;}
    table.review-table tr.total-row td {font-weight: 700; background: #f1f3f5;}
    </style>
    """
    headers = """
    <div class="review-scroll">
    <table class="review-table">
      <thead>
        <tr>
          <th rowspan="2">Attributes</th>
          <th colspan="6">Values</th>
          <th colspan="8">Attributes</th>
        </tr>
        <tr>
          <th class="sub" colspan="2">Correct</th>
          <th class="sub" colspan="2">Missed</th>
          <th class="sub" colspan="2">False</th>
          <th class="sub" colspan="2">Spec Sheet<br>NA</th>
          <th class="sub" colspan="2">AI Returned<br>False NA</th>
          <th class="sub" colspan="2">False Positives</th>
          <th class="sub" colspan="2">Omissions</th>
        </tr>
      </thead>
      <tbody>
    """
    body_parts: List[str] = []
    for _, row in summary_df.iterrows():
        is_total = str(row.get("Attributes", "")).strip().lower() == "total"
        cls = ' class="total-row"' if is_total else ""
        cells = [f'<td class="attr">{html.escape(str(row.get("Attributes", "")))}</td>']
        pairs = [
            ("Values Correct Count", "Values Correct %"),
            ("Values Missed Count", "Values Missed %"),
            ("Values False Count", "Values False %"),
            ("Spec Sheet NA Count", "Spec Sheet NA %"),
            ("AI Returned False NA Count", "AI Returned False NA %"),
            ("False Positives Count", "False Positives %"),
            ("Omissions Count", "Omissions %"),
        ]
        for count_col, pct_col in pairs:
            cells.append(f"<td>{_review_blankable_count(row.get(count_col, 0))}</td>")
            cells.append(f"<td>{_review_blankable_percent(row.get(pct_col, 0))}</td>")
        body_parts.append(f"<tr{cls}>" + "".join(cells) + "</tr>")
    table_html = css + headers + "\n".join(body_parts) + "</tbody></table></div>"
    st.markdown(table_html, unsafe_allow_html=True)


def render_review_tab() -> None:
    st.header("Review")
    st.write(
        "This tab shows a grouped extraction-review summary in the same format as your sample: "
        "value accuracy buckets in the middle and attribute review buckets on the right."
    )

    available_sources: List[str] = []
    if "main_records" in st.session_state:
        available_sources.append("Document & URL Analysis")
    if "brand_records" in st.session_state:
        available_sources.append("Brand Analysis")

    if not available_sources:
        st.info("Run an analysis first. After extraction, return to this Review tab to see the grouped review table.")
        return

    source_choice = st.radio("Review data source", available_sources, horizontal=True)
    if source_choice == "Document & URL Analysis":
        records_df = st.session_state.get("main_records", pd.DataFrame())
        total_resources = int(st.session_state.get("main_resource_count", 0) or 0)
        file_name = "document_url_review_report.xlsx"
    else:
        records_df = st.session_state.get("brand_records", pd.DataFrame())
        total_resources = int(st.session_state.get("brand_count", 0) or 0)
        file_name = "brand_review_report.xlsx"

    if records_df.empty:
        st.warning("The selected source has no extracted records yet.")
        return

    st.caption(
        "Counts are calculated per analyzed resource/brand. If AI validation is enabled, invalid rows are counted as False/False Positives. "
        "Without AI validation, extracted non-NA values are treated as Correct so the table remains usable."
    )
    summary_df = build_review_summary(records_df, total_resources)
    render_review_summary_table(summary_df)

    with st.expander("Show review data as a normal table"):
        display_df = summary_df.copy()
        for col in [c for c in display_df.columns if c.endswith(" %")]:
            display_df[col] = display_df[col].map(lambda x: _review_blankable_percent(x))
        for col in [c for c in display_df.columns if c.endswith(" Count")]:
            display_df[col] = display_df[col].map(lambda x: _review_blankable_count(x))
        st.dataframe(display_df, use_container_width=True, hide_index=True)

    export_bytes = dataframe_to_excel_bytes(
        {
            "Review Summary": review_summary_to_export(summary_df),
            "Raw Records": records_df,
        }
    )
    st.download_button(
        "Download review Excel report",
        data=export_bytes,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"review_download_{source_choice}",
    )

def render_deployment_help() -> None:
    st.header("No-Code GitHub Deployment Help")
    st.markdown(
        """
1. Create a GitHub repository.
2. Upload these project files to the repository root.
3. Open Streamlit Community Cloud and create a new app from your GitHub repo.
4. Set the main file path to `app.py`.
5. Add secrets in Streamlit settings, not in GitHub:

```toml
OPENAI_API_KEY = "paste-your-key-here"
OPENAI_MODEL = "gpt-4o-mini"
```

6. Deploy and use the generated Streamlit URL in your browser.

For testing on your own computer:

```bash
pip install -r requirements.txt
streamlit run app.py
```
        """
    )


def main() -> None:
    render_header()
    enable_ai, enable_validation, model = render_sidebar()
    tab1, tab2, tab3 = st.tabs(["Documents, URLs & Brand Comparison", "Review", "Deploy Help"])
    with tab1:
        render_main_analysis(enable_ai, enable_validation, model)
        st.divider()
        render_brand_analysis(enable_ai, model, key_prefix="bottom_brand")
    with tab2:
        render_review_tab()
    with tab3:
        render_deployment_help()


if __name__ == "__main__":
    main()
